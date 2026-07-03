"""
tools/file_generator.py
Generates PDF and DOCX report files from content.
"""

from langchain_core.tools import tool
import os
import re

OUTPUT_DIR = "./generated_files"


def _ensure_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


@tool
def generate_pdf(content: str, filename: str = "niarad_report") -> str:
    """
    Generate a PDF file from text content.
    Use this when the user asks to create a report, summary, or document
    as a downloadable PDF file.
    Input: content (the text to put in the PDF), filename (without extension).
    Output: path to the generated file.
    """
    _ensure_dir()
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.colors import HexColor

        safe_name = re.sub(r'[^\w\-]', '_', filename)
        filepath = os.path.join(OUTPUT_DIR, safe_name + ".pdf")

        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )

        styles = getSampleStyleSheet()
        story = []

        # Header
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Title'],
            fontSize=20,
            textColor=HexColor('#00FFAA'),
            spaceAfter=12
        )
        story.append(Paragraph("NIARAD REPORT", header_style))
        story.append(Spacer(1, 0.3*cm))

        # Content - split by lines
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=8
        )
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.2*cm))
            elif line.startswith('#'):
                heading = line.lstrip('#').strip()
                h_style = ParagraphStyle('H', parent=styles['Heading2'], fontSize=13, spaceAfter=6)
                story.append(Paragraph(heading, h_style))
            else:
                story.append(Paragraph(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))

        doc.build(story)
        return "PDF generated: " + filepath

    except ImportError:
        return "⚠ reportlab not installed. Run: pip install reportlab"
    except Exception as e:
        return "PDF generation failed: " + str(e)


@tool
def generate_docx(content: str, filename: str = "niarad_report") -> str:
    """
    Generate a Word DOCX file from text content.
    Use this when the user asks to create a Word document, report, or summary.
    Input: content (text to put in the document), filename (without extension).
    Output: path to the generated file.
    """
    _ensure_dir()
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        safe_name = re.sub(r'[^\w\-]', '_', filename)
        filepath = os.path.join(OUTPUT_DIR, safe_name + ".docx")

        doc = Document()

        # Title
        title = doc.add_heading('NIARAD REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Content
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        doc.save(filepath)
        return "DOCX generated: " + filepath

    except ImportError:
        return "⚠ python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return "DOCX generation failed: " + str(e)
